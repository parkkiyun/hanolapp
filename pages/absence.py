import streamlit as st
from app.auth_manager import AuthManager
from app.sidebar_manager import SidebarManager
from pathlib import Path
import os
from openpyxl.utils import get_column_letter
import openpyxl.cell.cell
from app.absence_excel_processing import process_excel

# 페이지 설정
st.set_page_config(
    page_title="결석신고서",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 권한 체크
auth_manager = AuthManager()
auth_manager.check_page_access("absence")

# 사이드바 렌더링
sidebar = SidebarManager()
sidebar.render_sidebar()

# 로그인 상태가 아니면 리다이렉트
if not st.session_state.get("authenticated", False):
    st.error("이 페이지는 교사 로그인이 필요합니다.")
    st.switch_page("pages/dashboard.py")

import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import tempfile
import base64
from openpyxl import load_workbook
import logging

# 배포 환경에서의 경로 처리
if os.getenv('STREAMLIT_SERVER_PATH'):  # Streamlit Cloud 환경인 경우
    ROOT_DIR = Path('/mount/src/hanolapp')  # Streamlit Cloud의 기본 경로
else:
    ROOT_DIR = Path(__file__).parent.parent.absolute()  # 로컬 환경

# 템플릿 디렉토리 경로 설정
TEMPLATE_DIR = ROOT_DIR / "templates"

# 템플릿 파일 경로 설정
TEMPLATE_FILES = {
    "출석인정결석": TEMPLATE_DIR / "출석인정 결석계 템플릿.docx",
    "질병결석": TEMPLATE_DIR / "질병결석계 템플릿.docx",
    "기타결석": TEMPLATE_DIR / "기타결석계 템플릿.docx",
}

# 템플릿 디렉토리 존재 여부 확인
if not TEMPLATE_DIR.exists():
    st.error(f"템플릿 디렉토리를 찾을 수 없습니다: {TEMPLATE_DIR}")
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

# 템플릿 파일 존재 여부 확인
for attendance_type, template_path in TEMPLATE_FILES.items():
    if not template_path.exists():
        st.error(f"템플릿 파일을 찾을 수 없습니다: {template_path}")

# 로고 파일 경로 수정
LOGO_PATH = ROOT_DIR / "images" / "logo.png"

# 로컬 이미지 파일을 Base64로 변환하는 함수
def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as img_file:
            encoded = base64.b64encode(img_file.read()).decode()
        return encoded
    except FileNotFoundError:
        st.error(f"로고 파일을 찾을 수 없습니다: {image_path}")
        return None

# Base64 인코딩된 로고 이미지
logo_base64 = get_base64_image(LOGO_PATH)

# 로고가 있을 때만 표시
if logo_base64:
    # 가운데 정렬을 위한 컬럼 설정
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        # 타이틀 가운데 정렬
        st.markdown("""
            <h1 style='text-align: center;'>스마트 결석신고서</h1>
        """, unsafe_allow_html=True)

        # 서브타이틀과 로고를 가운데 정렬
        st.markdown(f"""
            <div style="display: flex; align-items: center; justify-content: center; margin-bottom: 20px;">
                <img src="data:image/png;base64,{logo_base64}" alt="로고" style="margin-right: 10px; width: 30px; height: 30px;">
                <h3 style="margin: 0;">온양한올고등학교</h3>
            </div>
        """, unsafe_allow_html=True)
else:
    # 로고 없을 때도 가운데 정렬
    st.markdown("""
        <h1 style='text-align: center;'>스마트 결석신고서</h1>
        <h3 style='text-align: center;'>온양한올고등학교</h3>
    """, unsafe_allow_html=True)

st.markdown("---")


# Step 1: 정보 입력 단계
st.write("### 정보 입력")

if 'confirmation_date' not in st.session_state:
    st.session_state['confirmation_date'] = datetime.now()
if 'grade' not in st.session_state:
    st.session_state['grade'] = "1"
if 'class_name' not in st.session_state:
    st.session_state['class_name'] = "1"
if 'teacher_name' not in st.session_state:
    st.session_state['teacher_name'] = ""

# 학년, 반을 한 행에 배치
col1, col2 = st.columns(2)
with col1:
    grade = st.selectbox("학년", ["1학년", "2학년", "3학년"], key="grade_selectbox")[0]
with col2:
    class_name = st.selectbox("반", [f"{i}반" for i in range(1, 13)], key="class_selectbox").replace("반", "")

# 담임교사 성명, 결석확인일을 한 행에 배치
col3, col4 = st.columns(2)
with col3:
    teacher_name = st.text_input("담임교사 성명", key="teacher_name_input")
with col4:
    confirmation_date = st.date_input("결석확인일", st.session_state['confirmation_date'], key="confirmation_date_input")

if st.button("입력 완료", key="next_step_button"):
    st.session_state['confirmation_date_str'] = confirmation_date.strftime('%Y.%m.%d')
    st.session_state['grade'] = grade
    st.session_state['class_name'] = class_name
    st.session_state['teacher_name'] = teacher_name
    st.session_state['step'] = 2

# Step 2: 엑셀 파일 업로드 및 처리 단계
if 'step' in st.session_state and st.session_state['step'] == 2:
    st.write("### 엑셀 파일 업로드하기")

    uploaded_file = st.file_uploader("엑셀 파일 업로드", type=["xlsx", "xls"], key="file_uploader")

    if uploaded_file is not None:
        try:
            # 파일명 확인
            st.write("### 처리 중인 파일명:", uploaded_file.name)
            
            # 업로드된 파일을 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_path = temp_file.name

            # 새로운 process_excel 함수 사용
            data = process_excel(temp_path)
            
            # 임시 파일 삭제
            try:
                os.unlink(temp_path)
            except Exception as e:
                logging.warning(f"임시 파일 삭제 중 오류가 발생했습니다: {e}")

            # 결석확인일 추가
            data['결석확인일'] = st.session_state.get('confirmation_date_str', '')

            # 데이터 검증 및 표시
            if data.empty:
                st.error("처리할 데이터가 없습니다.")
            else:
                st.session_state['processed_data'] = data
                data = st.data_editor(data, key='data_editor', use_container_width=True)
                st.session_state['processed_data'] = data

                col1, col2 = st.columns([1, 1], gap="small")
                with col1:
                    st.button("이전 단계로 이동", on_click=lambda: st.session_state.update({'step': 1}))
                with col2:
                    st.button("DOCX 생성 및 다운로드하기", on_click=lambda: st.session_state.update({'step': 3}))

        except Exception as e:
            st.error(f"엑셀 데이터 처리 중 오류가 발생했습니다: {e}")
            logging.error(f"Excel processing error: {str(e)}")

# Step 3: DOCX 생성 및 다운로드 단계
if 'step' in st.session_state and st.session_state['step'] == 3:
    st.write("### DOCX 생성 및 다운로드")

    processed_data = st.session_state.get('processed_data', pd.DataFrame())

    if not processed_data.empty:
        st.write("처리된 데이터에 기반하여 DOCX 파일을 생성합니다.")

        for attendance_type, template_file_name in TEMPLATE_FILES.items():
            # 질병결석 케이스를 위한 특별 처리
            if attendance_type == '질병결석':
                filtered_data = processed_data[processed_data['출결구분'].str.contains('질병.*결석', regex=True)]
            else:
                filtered_data = processed_data[processed_data['출결구분'] == attendance_type]

            if filtered_data.empty:
                continue

            template_path = os.path.join(TEMPLATE_DIR, template_file_name)

            if not os.path.exists(template_path):
                st.error(f"템플릿 파일을 찾을 수 없습니다: {template_file_name}")
                continue

            output_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

            # 템플릿 파일 불러오기
            master_doc = Document(template_path)

            for idx, data_row in filtered_data.iterrows():
                # 각 학생의 데이터를 템플릿에 적용하여 새로운 문서에 추가
                temp_doc = Document(template_path)
                replacements = {
                    "{1}": str(st.session_state['grade']),
                    "{2}": str(st.session_state['class_name']),
                    "{3}": str(int(data_row['번호'])),
                    "{성명}": str(data_row['성명']),
                    "{결석사유}": str(data_row['사유']),
                    "{결석시작일}": str(data_row['결석시작일']),
                    "{결석종료일}": str(data_row['결석종료일']),
                    "{결석일수}": str(data_row['결석일수']),
                    "{결석확인일}": str(data_row['결석확인일']),
                    "{담임교사 성명}": str(st.session_state['teacher_name'])
                }

                # 텍스트 치환 함수 수정
                def replace_text_in_paragraph(paragraph, replacements):
                    for run in paragraph.runs:
                        text = run.text
                        for key, value in replacements.items():
                            if key in text:
                                text = text.replace(key, value)
                        run.text = text

                # 문서 내 텍스트 치환
                for paragraph in temp_doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

                # 표 안의 데이터 치환
                for table in temp_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, replacements)

                # 수정된 템플릿 내용을 통합 문서에 추가
                for element in temp_doc.element.body:
                    master_doc.element.body.append(element)

            # 마지막에 첫 번째 템플릿 페이지 제거
            if len(master_doc.element.body) > 1:
                master_doc.element.body.remove(master_doc.element.body[0])

            # 수정된 파일 저장
            master_doc.save(output_docx.name)

            # Streamlit에서 파일 다운로드 제공
            with open(output_docx.name, "rb") as f:
                st.download_button(
                    label=f"{st.session_state['grade']}학년 {st.session_state['class_name']}반 {attendance_type} 결석신고서({confirmation_date.month}월).docx 다운로드",
                    data=f,
                    file_name=f"{st.session_state['grade']}학년 {st.session_state['class_name']}반 {attendance_type} 결석신고서({confirmation_date.month}월).docx",
                    key=f"{attendance_type}_download"
                )
    else:
        st.error("처리된 데이터가 없습니다. 이전 단계에서 데이터를 확인해주세요.")


st.markdown("---")

st.subheader("프로그램 설명")
st.info(
    "스마트 결석신고서는 선생님들의 업무를 간소화하기 위해 설계되었습니다.\n\n"
    "엑셀파일 다운로드 경로:\n"
    "[나이스]-[출결현황 및 통계]-[학급별 출결현황]-[조회]-[엑셀 다운로드]\n\n"
    "오류나 작동 문제가 있으면 언제든지 말씀해주세요!"
)

st.markdown("---")
st.markdown("<div style='text-align: right;'>제작자: 박기윤</div>", unsafe_allow_html=True)

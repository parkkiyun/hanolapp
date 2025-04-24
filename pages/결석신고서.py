import streamlit as st
from app.sidebar_manager import SidebarManager
from pathlib import Path
import os
from openpyxl.utils import get_column_letter
import openpyxl.cell.cell
from app.absence_excel_processing import process_excel

# 페이지 설정
st.set_page_config(
    page_title="결석신고서",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://github.com/parkkiyun/hanolapp',
        'Report a bug': "https://github.com/parkkiyun/hanolapp/issues",
        'About': "# 스마트 문서 시스템 v1.0"
    }
)

# 사이드바 설정
st.markdown("""
    <style>
        [data-testid="stSidebar"] {
            min-width: 250px;
            max-width: 250px;
        }
    </style>
""", unsafe_allow_html=True)

# 사이드바 렌더링
sidebar = SidebarManager()
sidebar.render_sidebar()

import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import tempfile
import base64
from openpyxl import load_workbook
import logging

# 결석확인일 계산 함수 추가
def calculate_confirmation_date(end_date):
    """결석종료일로부터 5일 이내의 평일을 찾아 반환"""
    # datetime 객체로 변환
    if isinstance(end_date, str):
        try:
            end_date = datetime.strptime(end_date, '%Y-%m-%d')
        except ValueError:
            try:
                end_date = datetime.strptime(end_date, '%Y.%m.%d')
            except ValueError:
                return datetime.now().strftime('%Y.%m.%d')  # 기본값
    
    # 5일 이내의 날짜 중 평일(월-금) 찾기
    for i in range(1, 6):
        check_date = end_date + timedelta(days=i)
        # 평일이면 (0=월요일, 6=일요일)
        if check_date.weekday() < 5:
            return check_date.strftime('%Y.%m.%d')
    
    # 5일 이내에 평일이 없으면 종료일로부터 다음 월요일
    while (end_date + timedelta(days=1)).weekday() != 0:
        end_date = end_date + timedelta(days=1)
    return (end_date + timedelta(days=1)).strftime('%Y.%m.%d')

# 배포 환경에서의 경로 처리
if os.getenv('STREAMLIT_SERVER_PATH'):  # Streamlit Cloud 환경인 경우
    ROOT_DIR = Path('/mount/src/hanolapp')  # Streamlit Cloud의 기본 경로
else:
    ROOT_DIR = Path(__file__).parent.parent.absolute()  # 로컬 환경

# 템플릿 디렉토리 경로 설정
TEMPLATE_DIR = ROOT_DIR / "templates"

# 템플릿 파일 경로 설정
TEMPLATE_FILES = {
    "출석인정결석": TEMPLATE_DIR / "출석인정 결석계 템플릿.docx",
    "질병결석": TEMPLATE_DIR / "질병결석계 템플릿.docx",
    "기타결석": TEMPLATE_DIR / "기타결석계 템플릿.docx",
}

# 템플릿 디렉토리 존재 여부 확인
if not TEMPLATE_DIR.exists():
    st.error(f"템플릿 디렉토리를 찾을 수 없습니다: {TEMPLATE_DIR}")
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

# 템플릿 파일 존재 여부 확인
for attendance_type, template_path in TEMPLATE_FILES.items():
    if not template_path.exists():
        st.error(f"템플릿 파일을 찾을 수 없습니다: {template_path}")

# 로고 파일 경로 수정
LOGO_PATH = ROOT_DIR / "images" / "logo.png"

# 로컬 이미지 파일을 Base64로 변환하는 함수
def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as img_file:
            encoded = base64.b64encode(img_file.read()).decode()
        return encoded
    except FileNotFoundError:
        st.error(f"로고 파일을 찾을 수 없습니다: {image_path}")
        return None

# Base64 인코딩된 로고 이미지
logo_base64 = get_base64_image(LOGO_PATH)

# 로고가 있을 때만 표시
if logo_base64:
    # 가운데 정렬을 위한 컬럼 설정
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        # 타이틀 가운데 정렬
        st.markdown("""
            <h1 style='text-align: center;'>스마트 결석신고서</h1>
        """, unsafe_allow_html=True)

        # 서브타이틀과 로고를 가운데 정렬
        st.markdown(f"""
            <div style="display: flex; align-items: center; justify-content: center; margin-bottom: 20px;">
                <img src="data:image/png;base64,{logo_base64}" alt="로고" style="margin-right: 10px; width: 30px; height: 30px;">
                <h3 style="margin: 0;">한올고등학교</h3>
            </div>
        """, unsafe_allow_html=True)
else:
    # 로고 없을 때도 가운데 정렬
    st.markdown("""
        <h1 style='text-align: center;'>스마트 결석신고서</h1>
        <h3 style='text-align: center;'>한올고등학교</h3>
    """, unsafe_allow_html=True)

st.markdown("---")


# Step 1: 정보 입력 단계
st.write("### 정보 입력")

if 'grade' not in st.session_state:
    st.session_state['grade'] = "1"
if 'class_name' not in st.session_state:
    st.session_state['class_name'] = "1"
if 'teacher_name' not in st.session_state:
    st.session_state['teacher_name'] = ""

# 학년, 반을 한 행에 배치
col1, col2 = st.columns(2)
with col1:
    grade = st.selectbox("학년", ["1학년", "2학년", "3학년"], key="grade_selectbox")[0]
with col2:
    class_name = st.selectbox("반", [f"{i}반" for i in range(1, 13)], key="class_selectbox").replace("반", "")

# 담임교사 성명만 입력 (결석확인일은 자동 계산)
teacher_name = st.text_input("담임교사 성명", key="teacher_name_input")

if st.button("입력 완료", key="next_step_button"):
    st.session_state['grade'] = grade
    st.session_state['class_name'] = class_name
    st.session_state['teacher_name'] = teacher_name
    st.session_state['step'] = 2

# Step 2: 엑셀 파일 업로드 및 행 선택 단계
if 'step' in st.session_state and st.session_state['step'] == 2:
    st.write("### 엑셀 파일 업로드 및 행 선택")

    uploaded_file = st.file_uploader("엑셀 파일 업로드", type=["xlsx", "xls"], key="file_uploader")

    if uploaded_file is not None:
        try:
            # 업로드된 파일을 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_path = temp_file.name

            # 새로운 process_excel 함수 사용
            data = process_excel(temp_path)
            
            # 임시 파일 삭제
            try:
                os.unlink(temp_path)
            except Exception as e:
                logging.warning(f"임시 파일 삭제 중 오류가 발생했습니다: {e}")

            # 출결구분 필터링 추가 - 질병결석, 출석인정결석, 기타결석만 유지
            valid_attendance_types = ["질병결석", "출석인정결석", "기타결석"]
            data = data[data['출결구분'].isin(valid_attendance_types)]

            # 각 학생의 결석종료일에 따라 결석확인일 자동 계산
            data['결석확인일'] = data['결석종료일'].apply(calculate_confirmation_date)

            # 데이터 검증 및 표시
            if data.empty:
                st.error("처리할 데이터가 없습니다. 질병결석, 출석인정결석, 기타결석 데이터가 있는지 확인하세요.")
            else:
                st.session_state['processed_data'] = data
                
                # 각 출결구분별로 데이터 분리
                attendance_types = data['출결구분'].unique()
                st.write("### 결석 신고서 선택")
                st.info("결석신고서를 생성할 학생을 선택해주세요.")
                
                all_selected_indices = []  # 모든 선택된 인덱스를 저장할 리스트
                
                # 탭 생성
                tabs = st.tabs([f"{attendance_type}" for attendance_type in attendance_types])
                
                for tab, attendance_type in zip(tabs, attendance_types):
                    with tab:
                        type_data = data[data['출결구분'] == attendance_type].copy()
                        # '선택' 컬럼 추가
                        type_data['선택'] = False
                        
                        # 모두 선택 버튼 추가
                        col1, col2 = st.columns([1, 4])
                        with col1:
                            if st.button("모두 선택", key=f"select_all_{attendance_type}"):
                                type_data['선택'] = True
                        
                        # 데이터프레임을 표시하고 선택 가능하게 만들기
                        edited_df = st.data_editor(
                            type_data,
                            hide_index=True,
                            column_config={
                                "선택": st.column_config.CheckboxColumn(
                                    "선택",
                                    help="DOCX 파일로 추출할 행을 선택하세요",
                                    default=False,
                                )
                            },
                            key=f"editor_{attendance_type}"
                        )
                        
                        # 선택된 행의 인덱스 수집
                        selected_rows = edited_df[edited_df["선택"]]
                        all_selected_indices.extend(selected_rows.index.tolist())
                
                # 선택된 행 저장
                if all_selected_indices:
                    selected_data = data.loc[all_selected_indices]
                    st.session_state['selected_data'] = selected_data
                    
                    # 선택된 데이터 요약 표시
                    st.write("### 선택된 데이터 요약")
                    summary_cols = st.columns(len(attendance_types))
                    for col, attendance_type in zip(summary_cols, attendance_types):
                        with col:
                            count = len(selected_data[selected_data['출결구분'] == attendance_type])
                            st.metric(f"{attendance_type}", f"{count}건")
                else:
                    st.session_state['selected_data'] = pd.DataFrame()
                
                # 버튼 배치
                col1, col2, col3 = st.columns([1, 1, 1], gap="small")
                with col1:
                    st.button("이전 단계로 이동", on_click=lambda: st.session_state.update({'step': 1}))
                with col2:
                    st.button("선택 초기화", on_click=lambda: st.session_state.update({'selected_data': pd.DataFrame()}))
                with col3:
                    st.button("DOCX 생성 및 다운로드하기", on_click=lambda: st.session_state.update({'step': 3}))

        except Exception as e:
            st.error(f"엑셀 데이터 처리 중 오류가 발생했습니다: {e}")
            logging.error(f"Excel processing error: {str(e)}")

# Step 3: DOCX 생성 및 다운로드 단계
if 'step' in st.session_state and st.session_state['step'] == 3:
    st.write("### DOCX 생성 및 다운로드")

    processed_data = st.session_state.get('selected_data', pd.DataFrame())

    if not processed_data.empty:
        for attendance_type, template_file_name in TEMPLATE_FILES.items():
            # 모든 출결구분에 대해 동일한 필터링 적용
            filtered_data = processed_data[processed_data['출결구분'] == attendance_type]

            if filtered_data.empty:
                continue

            template_path = os.path.join(TEMPLATE_DIR, template_file_name)

            if not os.path.exists(template_path):
                st.error(f"템플릿 파일을 찾을 수 없습니다: {template_file_name}")
                continue

            output_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

            # 템플릿 파일 불러오기
            master_doc = Document(template_path)

            for idx, data_row in filtered_data.iterrows():
                # 각 학생의 데이터를 템플릿에 적용하여 새로운 문서에 추가
                temp_doc = Document(template_path)
                replacements = {
                    "{1}": str(st.session_state['grade']),
                    "{2}": str(st.session_state['class_name']),
                    "{3}": str(int(data_row['번호'])),
                    "{성명}": str(data_row['성명']),
                    "{결석사유}": str(data_row['사유']),
                    "{결석시작일}": str(data_row['결석시작일']),
                    "{결석종료일}": str(data_row['결석종료일']),
                    "{결석일수}": str(data_row['결석일수']),
                    "{결석확인일}": str(data_row['결석확인일']),
                    "{담임교사 성명}": str(st.session_state['teacher_name'])
                }

                # 텍스트 치환 함수 수정
                def replace_text_in_paragraph(paragraph, replacements):
                    for run in paragraph.runs:
                        text = run.text
                        for key, value in replacements.items():
                            if key in text:
                                text = text.replace(key, value)
                        run.text = text

                # 문서 내 텍스트 치환
                for paragraph in temp_doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

                # 표 안의 데이터 치환
                for table in temp_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, replacements)

                # 수정된 템플릿 내용을 통합 문서에 추가
                for element in temp_doc.element.body:
                    master_doc.element.body.append(element)

            # 마지막에 첫 번째 템플릿 페이지 제거
            if len(master_doc.element.body) > 1:
                master_doc.element.body.remove(master_doc.element.body[0])

            # 수정된 파일 저장
            master_doc.save(output_docx.name)

            # 다운로드 버튼 레이블에 현재 달을 표시
            current_month = datetime.now().month
            
            # Streamlit에서 파일 다운로드 제공
            with open(output_docx.name, "rb") as f:
                st.download_button(
                    label=f"{st.session_state['grade']}학년 {st.session_state['class_name']}반 {attendance_type} 결석신고서({current_month}월).docx 다운로드",
                    data=f,
                    file_name=f"{st.session_state['grade']}학년 {st.session_state['class_name']}반 {attendance_type} 결석신고서({current_month}월).docx",
                    key=f"{attendance_type}_download"
                )
    else:
        st.error("처리된 데이터가 없습니다. 이전 단계에서 데이터를 확인해주세요.")


st.markdown("---")

st.subheader("프로그램 설명")
st.info(
    "스마트 결석신고서는 선생님들의 업무를 간소화하기 위해 설계되었습니다.\n\n"
    "엑셀파일 다운로드 경로:\n"
    "[나이스]-[출결현황 및 통계]-[학급별 출결현황]-[조회]-[엑셀 다운로드]\n\n"
    "오류나 작동 문제가 있으면 언제든지 말씀해주세요!"
)

st.markdown("---")
st.markdown("<div style='text-align: right;'>제작자: 박기윤</div>", unsafe_allow_html=True)
